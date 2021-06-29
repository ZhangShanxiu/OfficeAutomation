#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2021/6/18 17:05
# @Author  : Zhang Shanxiu
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #对齐样式
from docx.shared import Pt  #字体大小
from docx.oxml.ns import qn #中文格式
from docx.shared import Inches #图片尺寸
import time


def main():
    price = input('请输入今日黄金价格：')
    customers = ['客户1', '客户2', '客户3', '客户4', '客户5', '客户6', '客户7', '客户8', '客户9']
    today = time.strftime('%Y-%m-%d', time.localtime())
    today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y = '年', m = '月', d = '日')
    print(today)

    for customer in customers:
        document = Document()
        document.styles['Normal'].font.name=u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_picture('./jpgs/2.JPG', width=Inches(2.0))
        document.add_picture('./jpgs/3.JPG', width=Inches(2.0))

        p1 = document.add_paragraph()
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = p1.add_run('关于下达 %s 价格通知' % (today))
        run1.font.name = '微软雅黑'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        run1.font.size = Pt(21)
        run1.font.bold = True

        p2 = document.add_paragraph()
        # p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = p2.add_run(customer + ':')
        run2.font.name = '仿宋_GB2312'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run2.font.size = Pt(21)
        run2.font.bold = True

        p3 = document.add_paragraph()
        # p3.alignment = WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY
        run3 = p3.add_run('        根据公司的安排，为提供优质的客户服务，我单位拟定了今日黄金价格为 %s 元，特此通知。' % price)
        run3.font.name = '仿宋_GB2312'
        run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run3.font.size = Pt(21)
        run3.font.bold = True

        p4 = document.add_paragraph()
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run4 = p4.add_run('联系人：红心石，电话：15817200324')
        run4.font.name = '仿宋_GB2312'
        run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run4.font.size = Pt(21)
        run4.font.bold = True

        p1.space_after = Pt(5)
        p1.space_before = Pt(5)

        document.save('./docs/%s-价格通知.doc' % customer)


if __name__ == '__main__':
    main()